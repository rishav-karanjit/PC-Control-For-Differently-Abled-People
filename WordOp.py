import docx
import os
class Word:
    def __init__(self):
        self.mydoc = docx.Document()
    def WReadAll(self,Doc_Name):
        Doc_Name += ".docx"
        doc = docx.Document(os.path.join(".\Doc",Doc_Name))
        all_para_list = doc.paragraphs
        all_para = ""
        for para in all_para_list:
            all_para = all_para + "\n" + para.text
        return all_para
    
    def WReadPara(self,Doc_Name,ParaNum):
        Doc_Name += ".docx"
        doc = docx.Document(os.path.join(".\Doc",Doc_Name))
        ParaNum-=1
        para = doc.paragraphs[ParaNum]
        while(para.text == ""):
            ParaNum+=1
            para = doc.paragraphs[ParaNum]
        return para.text

    def WAddPara(self,text):
        self.mydoc.add_paragraph(text)

    def WAddHeading(self,text,lvl):
        self.mydoc.add_heading(text,lvl)

    def SaveWrittenTxt(self,Doc_Name):
        Doc_Name += ".docx"
        self.mydoc.save(os.path.join(".\Doc",Doc_Name))
        self.mydoc = docx.Document()

# W = Word()
# W.WAddHeading("Heading 0", 0)
# W.WAddPara("ABCD")
# W.WAddPara("XYZ")
# W.WAddHeading("Heading 1", 1)
# W.WAddPara("QWERTY")
# W.WAddPara("UIOP")
# W.WAddHeading("Heading 2", 2)
# W.SaveWrittenTxt("ABC")